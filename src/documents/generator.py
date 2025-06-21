"""
Document Generator - Create formatted documents with video summaries
Supports TXT, DOCX, and PDF formats
"""

import asyncio
import logging
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, Callable

# Document generation libraries
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

from config.settings import get_settings
from utils.validators import sanitize_filename

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentGenerator:
    """Generate formatted documents with video summaries."""
    
    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "yt_summarizer_docs"
        self.temp_dir.mkdir(exist_ok=True)
        self._register_unicode_fonts()
    
    def _register_unicode_fonts(self):
        """Register Unicode fonts for PDF generation to support Cyrillic characters."""
        try:
            # Define font paths - try to find suitable Unicode fonts
            font_paths = [
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
                '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
                '/System/Library/Fonts/Arial.ttf',  # macOS
                'C:\\Windows\\Fonts\\arial.ttf',    # Windows
            ]
            
            fonts_registered = False
            
            # Try to register DejaVu fonts (best for Cyrillic support)
            try:
                dejavu_regular = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
                dejavu_bold = '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'
                
                if Path(dejavu_regular).exists() and Path(dejavu_bold).exists():
                    pdfmetrics.registerFont(TTFont('DejaVuSans', dejavu_regular))
                    pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', dejavu_bold))
                    
                    # Add font mapping for the family
                    addMapping('DejaVuSans', 0, 0, 'DejaVuSans')       # normal
                    addMapping('DejaVuSans', 0, 1, 'DejaVuSans')       # italic (use normal)
                    addMapping('DejaVuSans', 1, 0, 'DejaVuSans-Bold')  # bold
                    addMapping('DejaVuSans', 1, 1, 'DejaVuSans-Bold')  # bold italic (use bold)
                    
                    self.unicode_font_family = 'DejaVuSans'
                    fonts_registered = True
                    logger.info("Registered DejaVu fonts for PDF generation")
                    
            except Exception as e:
                logger.warning(f"Could not register DejaVu fonts: {e}")
            
            # Fallback to Liberation fonts
            if not fonts_registered:
                try:
                    liberation_regular = '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf'
                    liberation_bold = '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf'
                    
                    if Path(liberation_regular).exists() and Path(liberation_bold).exists():
                        pdfmetrics.registerFont(TTFont('LiberationSans', liberation_regular))
                        pdfmetrics.registerFont(TTFont('LiberationSans-Bold', liberation_bold))
                        
                        addMapping('LiberationSans', 0, 0, 'LiberationSans')
                        addMapping('LiberationSans', 0, 1, 'LiberationSans')
                        addMapping('LiberationSans', 1, 0, 'LiberationSans-Bold')
                        addMapping('LiberationSans', 1, 1, 'LiberationSans-Bold')
                        
                        self.unicode_font_family = 'LiberationSans'
                        fonts_registered = True
                        logger.info("Registered Liberation fonts for PDF generation")
                        
                except Exception as e:
                    logger.warning(f"Could not register Liberation fonts: {e}")
            
            # Final fallback - use default fonts (may not support Cyrillic properly)
            if not fonts_registered:
                self.unicode_font_family = 'Helvetica'
                logger.warning("Using default fonts - Cyrillic characters may not display properly")
                
        except Exception as e:
            logger.error(f"Error registering Unicode fonts: {e}")
            self.unicode_font_family = 'Helvetica'

    async def _run_in_executor(self, func: Callable) -> None:
        """Run a function in executor to avoid blocking."""
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, func)

    def _get_pdf_styles(self):
        """Get common PDF styles for document generation."""
        styles = getSampleStyleSheet()
        
        return {
            'title': ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=self.unicode_font_family,
                fontSize=18,
                alignment=1,  # Center
                spaceAfter=30
            ),
            'heading': ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName=self.unicode_font_family,
                fontSize=14,
                textColor='#2E74B5',
                spaceAfter=12
            ),
            'normal': ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=self.unicode_font_family,
                fontSize=11,
                spaceAfter=6
            ),
            'footer': ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontName=self.unicode_font_family,
                fontSize=10,
                alignment=1,  # Center
                textColor='#666666'
            )
        }

    def _generate_filename(self, title: str, doc_type: str, output_format: str, corrected: bool = False) -> str:
        """Generate standardized filename."""
        safe_title = sanitize_filename(title)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if corrected:
            return f"{safe_title}_corrected_{doc_type}_{timestamp}.{output_format}"
        else:
            return f"{safe_title}_{doc_type}_{timestamp}.{output_format}"

    def _add_pdf_footer(self, story: list, pdf_styles: dict, footer_text: str):
        """Add footer to PDF document."""
        story.append(Spacer(1, 40))
        story.append(Paragraph(footer_text, pdf_styles['footer']))

    async def create_document(
        self, 
        video_data: Dict, 
        summary_data: Dict, 
        output_format: str = 'txt'
    ) -> Optional[Path]:
        """
        Create formatted document with video summary.
        
        Args:
            video_data: Processed video data
            summary_data: AI-generated summary data
            output_format: Output format ('txt', 'docx', 'pdf')
            
        Returns:
            Path to generated document or None on error
        """
        try:
            if output_format not in settings.supported_formats:
                raise ValueError(f"Unsupported format: {output_format}")
            
            # Prepare document content
            content = self._prepare_document_content(video_data, summary_data)
            
            # Generate filename
            video_title = video_data.get('video_info', {}).get('title', 'Unknown Video')
            filename = self._generate_filename(video_title, 'summary', output_format)
            document_path = self.temp_dir / filename
            
            # Create document based on format
            if output_format == 'txt':
                document_path = await self._create_txt_document(content, document_path)
            elif output_format == 'docx':
                document_path = await self._create_docx_document(content, document_path)
            elif output_format == 'pdf':
                document_path = await self._create_pdf_document(content, document_path)
            
            logger.info(f"Created {output_format.upper()} document: {document_path}")
            return document_path
            
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            return None

    def _prepare_document_content(self, video_data: Dict, summary_data: Dict) -> Dict:
        """Prepare structured content for document generation."""
        video_info = video_data.get('video_info', {})
        transcripts = video_data.get('transcripts', {})
        summary = summary_data.get('summary', {})
        
        # Format duration
        duration = video_info.get('duration', 0)
        duration_str = f"{duration // 60}:{duration % 60:02d}"
        
        # Format date
        upload_date = video_info.get('upload_date', '')
        if upload_date and len(upload_date) >= 8:
            try:
                formatted_date = f"{upload_date[6:8]}/{upload_date[4:6]}/{upload_date[:4]}"
            except:
                formatted_date = upload_date
        else:
            formatted_date = upload_date or 'Unknown'
        
        return {
            'title': f"YouTube Video Summary: {video_info.get('title', 'Unknown Title')}",
            'video_info': {
                'title': video_info.get('title', 'Unknown Title'),
                'uploader': video_info.get('uploader', 'Unknown'),
                'duration': duration_str,
                'upload_date': formatted_date,
                'video_id': video_info.get('id', ''),
                'view_count': f"{video_info.get('view_count', 0):,}" if video_info.get('view_count') else 'Unknown'
            },
            'summary': {
                'executive_summary': summary.get('executive_summary', 'No executive summary available.'),
                'key_points': summary.get('key_points', []),
                'detailed_summary': summary.get('detailed_summary', 'No detailed summary available.'),
                'timestamps': summary.get('timestamps', []),
                'takeaways': summary.get('takeaways', [])
            },
            'transcript_info': {
                'available': bool(transcripts.get('transcripts')),
                'language': transcripts.get('language', 'Unknown'),
                'is_generated': transcripts.get('is_generated', None),
                'total_segments': transcripts.get('total_segments', 0)
            },
            'ai_info': {
                'model': summary_data.get('ai_model', settings.openai_model),
                'tokens_used': summary_data.get('tokens_used', 0),
                'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')
            }
        }

    def _detect_content_type(self, content: Dict) -> str:
        """Detect content type based on structure."""
        if 'subtitles_text' in content:
            return 'subtitles'
        elif 'summary' in content:
            return 'summary'
        else:
            return 'unknown'

    async def _create_txt_document(self, content: Dict, output_path: Path) -> Path:
        """Create plain text document for both summaries and subtitles."""
        def generate_txt():
            lines = []
            content_type = self._detect_content_type(content)
            
            # Header
            lines.append("=" * 80)
            lines.append(content['title'])
            lines.append("=" * 80)
            lines.append("")
            
            if content_type == 'summary':
                # Video Information
                lines.append("VIDEO INFORMATION")
                lines.append("-" * 40)
                lines.append(f"Title: {content['video_info']['title']}")
                lines.append(f"Channel: {content['video_info']['uploader']}")
                lines.append(f"Duration: {content['video_info']['duration']}")
                lines.append(f"Upload Date: {content['video_info']['upload_date']}")
                lines.append(f"Video ID: {content['video_info']['video_id']}")
                lines.append(f"View Count: {content['video_info']['view_count']}")
                lines.append("")
                
                # Executive Summary
                lines.append("EXECUTIVE SUMMARY")
                lines.append("-" * 40)
                lines.append(content['summary']['executive_summary'])
                lines.append("")
                
                # Key Points
                if content['summary']['key_points']:
                    lines.append("KEY POINTS")
                    lines.append("-" * 40)
                    for i, point in enumerate(content['summary']['key_points'], 1):
                        lines.append(f"{i}. {point}")
                    lines.append("")
                
                # Detailed Summary
                if content['summary']['detailed_summary']:
                    lines.append("DETAILED SUMMARY")
                    lines.append("-" * 40)
                    lines.append(content['summary']['detailed_summary'])
                    lines.append("")
                
                # Important Timestamps
                if content['summary']['timestamps']:
                    lines.append("IMPORTANT TIMESTAMPS")
                    lines.append("-" * 40)
                    for timestamp in content['summary']['timestamps']:
                        lines.append(f"• {timestamp}")
                    lines.append("")
                
                # Takeaways
                if content['summary']['takeaways']:
                    lines.append("KEY TAKEAWAYS")
                    lines.append("-" * 40)
                    for takeaway in content['summary']['takeaways']:
                        lines.append(f"• {takeaway}")
                    lines.append("")
                
                # Transcript Information
                lines.append("TRANSCRIPT INFORMATION")
                lines.append("-" * 40)
                lines.append(f"Available: {'Yes' if content['transcript_info']['available'] else 'No'}")
                if content['transcript_info']['available']:
                    lines.append(f"Language: {content['transcript_info']['language']}")
                    lines.append(f"Type: {'Auto-generated' if content['transcript_info']['is_generated'] else 'Manual'}")
                    lines.append(f"Total Segments: {content['transcript_info']['total_segments']}")
                lines.append("")
                
                # AI Generation Info
                lines.append("SUMMARY GENERATION INFO")
                lines.append("-" * 40)
                lines.append(f"AI Model: {content['ai_info']['model']}")
                lines.append(f"Tokens Used: {content['ai_info']['tokens_used']}")
                lines.append(f"Generated: {content['ai_info']['generated_at']}")
                lines.append("")
                
                # Footer
                lines.append("=" * 80)
                lines.append(f"Generated by {settings.bot_name} v{settings.bot_version}")
                lines.append("=" * 80)
                
            elif content_type == 'subtitles':
                # Video Information
                lines.append("ИНФОРМАЦИЯ О ВИДЕО")
                lines.append("-" * 40)
                lines.append(f"Название: {content['video_info']['title']}")
                lines.append(f"Канал: {content['video_info']['channel']}")
                lines.append(f"Длительность: {content['video_info']['duration']}")
                lines.append(f"ID видео: {content['video_info']['video_id']}")
                lines.append(f"Язык: {content['video_info']['language']} ({content['video_info']['language_code']})")
                lines.append(f"Тип: {'Автогенерированные' if content['video_info']['auto_generated'] else 'Ручные'}")
                
                if content['video_info']['corrected']:
                    lines.append(f"Обработка: ИИ коррекция ({content['video_info']['correction_method']})")
                
                lines.append(f"Количество сегментов: {content['video_info']['subtitle_count']}")
                lines.append("")
                
                # Subtitles
                lines.append("СУБТИТРЫ")
                lines.append("-" * 40)
                lines.append(content['subtitles_text'])
                lines.append("")
                
                # Footer
                lines.append("=" * 80)
                lines.append(f"Сгенерировано: {content['generation_info']['generated_at']}")
                lines.append(f"Тип документа: {content['generation_info']['format']}")
                lines.append(f"Создано с помощью {settings.bot_name} v{settings.bot_version}")
                lines.append("=" * 80)
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(lines))
        
        await self._run_in_executor(generate_txt)
        return output_path

    async def _create_docx_document(self, content: Dict, output_path: Path) -> Path:
        """Create Microsoft Word document for both summaries and subtitles."""
        def generate_docx():
            doc = Document()
            content_type = self._detect_content_type(content)
            
            # Title
            title = doc.add_heading(content['title'], 0)
            title.alignment = 1  # Center alignment
            
            if content_type == 'summary':
                # Video Information Section
                doc.add_heading('Video Information', level=1)
                
                # Create a table for video info
                table = doc.add_table(rows=6, cols=2)
                table.style = 'Table Grid'
                
                info_items = [
                    ('Title', content['video_info']['title']),
                    ('Channel', content['video_info']['uploader']),
                    ('Duration', content['video_info']['duration']),
                    ('Upload Date', content['video_info']['upload_date']),
                    ('Video ID', content['video_info']['video_id']),
                    ('View Count', content['video_info']['view_count'])
                ]
                
                for i, (label, value) in enumerate(info_items):
                    table.cell(i, 0).text = label
                    table.cell(i, 1).text = str(value)
                
                # Executive Summary
                doc.add_heading('Executive Summary', level=1)
                doc.add_paragraph(content['summary']['executive_summary'])
                
                # Key Points
                if content['summary']['key_points']:
                    doc.add_heading('Key Points', level=1)
                    for point in content['summary']['key_points']:
                        doc.add_paragraph(point, style='List Bullet')
                
                # Detailed Summary
                if content['summary']['detailed_summary']:
                    doc.add_heading('Detailed Summary', level=1)
                    doc.add_paragraph(content['summary']['detailed_summary'])
                
                # Important Timestamps
                if content['summary']['timestamps']:
                    doc.add_heading('Important Timestamps', level=1)
                    for timestamp in content['summary']['timestamps']:
                        doc.add_paragraph(timestamp, style='List Bullet')
                
                # Key Takeaways
                if content['summary']['takeaways']:
                    doc.add_heading('Key Takeaways', level=1)
                    for takeaway in content['summary']['takeaways']:
                        doc.add_paragraph(takeaway, style='List Bullet')
                
                # Technical Information
                doc.add_heading('Technical Information', level=1)
                
                tech_info = doc.add_paragraph()
                tech_info.add_run('Transcript Available: ').bold = True
                tech_info.add_run('Yes' if content['transcript_info']['available'] else 'No')
                tech_info.add_run('\n')
                
                if content['transcript_info']['available']:
                    tech_info.add_run('Language: ').bold = True
                    tech_info.add_run(content['transcript_info']['language'])
                    tech_info.add_run('\n')
                    tech_info.add_run('Type: ').bold = True
                    tech_info.add_run('Auto-generated' if content['transcript_info']['is_generated'] else 'Manual')
                    tech_info.add_run('\n')
                
                tech_info.add_run('AI Model: ').bold = True
                tech_info.add_run(content['ai_info']['model'])
                tech_info.add_run('\n')
                tech_info.add_run('Tokens Used: ').bold = True
                tech_info.add_run(str(content['ai_info']['tokens_used']))
                tech_info.add_run('\n')
                tech_info.add_run('Generated: ').bold = True
                tech_info.add_run(content['ai_info']['generated_at'])
                
                # Footer
                doc.add_page_break()
                footer_para = doc.add_paragraph()
                footer_para.alignment = 1  # Center alignment
                footer_run = footer_para.add_run(f"Generated by {settings.bot_name} v{settings.bot_version}")
                footer_run.italic = True
                
            elif content_type == 'subtitles':
                # Video Information Section
                doc.add_heading('Информация о видео', level=1)
                
                # Create a table for video info
                table = doc.add_table(rows=7, cols=2)
                table.style = 'Table Grid'
                
                info_items = [
                    ('Название', content['video_info']['title']),
                    ('Канал', content['video_info']['channel']),
                    ('Длительность', content['video_info']['duration']),
                    ('ID видео', content['video_info']['video_id']),
                    ('Язык', f"{content['video_info']['language']} ({content['video_info']['language_code']})"),
                    ('Тип', 'Автогенерированные' if content['video_info']['auto_generated'] else 'Ручные'),
                    ('Количество сегментов', str(content['video_info']['subtitle_count']))
                ]
                
                for i, (label, value) in enumerate(info_items):
                    table.cell(i, 0).text = label
                    table.cell(i, 1).text = str(value)
                
                if content['video_info']['corrected']:
                    doc.add_paragraph(f"✨ Обработано ИИ: {content['video_info']['correction_method']}")
                
                # Subtitles Section
                doc.add_heading('Субтитры', level=1)
                
                # Split subtitles into paragraphs for better readability
                subtitle_lines = content['subtitles_text'].split('\n')
                for line in subtitle_lines:
                    if line.strip():
                        doc.add_paragraph(line.strip())
                
                # Technical Information
                doc.add_heading('Информация о создании', level=1)
                
                tech_info = doc.add_paragraph()
                tech_info.add_run('Создано: ').bold = True
                tech_info.add_run(content['generation_info']['generated_at'])
                tech_info.add_run('\n')
                tech_info.add_run('Тип документа: ').bold = True
                tech_info.add_run(content['generation_info']['format'])
                
                # Footer
                doc.add_page_break()
                footer_para = doc.add_paragraph()
                footer_para.alignment = 1  # Center alignment
                footer_run = footer_para.add_run(f"Создано с помощью {settings.bot_name} v{settings.bot_version}")
                footer_run.italic = True
            
            # Save document
            doc.save(output_path)
        
        await self._run_in_executor(generate_docx)
        return output_path

    async def _create_pdf_document(self, content: Dict, output_path: Path) -> Path:
        """Create PDF document for both summaries and subtitles."""
        def generate_pdf():
            doc = SimpleDocTemplate(str(output_path), pagesize=letter)
            story = []
            pdf_styles = self._get_pdf_styles()
            content_type = self._detect_content_type(content)
            
            # Title
            story.append(Paragraph(content['title'], pdf_styles['title']))
            story.append(Spacer(1, 20))
            
            if content_type == 'summary':
                # Video Information
                story.append(Paragraph('Video Information', pdf_styles['heading']))
                
                video_info_text = f"""
                <b>Title:</b> {content['video_info']['title']}<br/>
                <b>Channel:</b> {content['video_info']['uploader']}<br/>
                <b>Duration:</b> {content['video_info']['duration']}<br/>
                <b>Upload Date:</b> {content['video_info']['upload_date']}<br/>
                <b>Video ID:</b> {content['video_info']['video_id']}<br/>
                <b>View Count:</b> {content['video_info']['view_count']}
                """
                story.append(Paragraph(video_info_text, pdf_styles['normal']))
                story.append(Spacer(1, 20))
                
                # Executive Summary
                story.append(Paragraph('Executive Summary', pdf_styles['heading']))
                story.append(Paragraph(content['summary']['executive_summary'], pdf_styles['normal']))
                story.append(Spacer(1, 20))
                
                # Key Points
                if content['summary']['key_points']:
                    story.append(Paragraph('Key Points', pdf_styles['heading']))
                    for point in content['summary']['key_points']:
                        story.append(Paragraph(f"• {point}", pdf_styles['normal']))
                    story.append(Spacer(1, 20))
                
                # Detailed Summary
                if content['summary']['detailed_summary']:
                    story.append(Paragraph('Detailed Summary', pdf_styles['heading']))
                    story.append(Paragraph(content['summary']['detailed_summary'], pdf_styles['normal']))
                    story.append(Spacer(1, 20))
                
                # Important Timestamps
                if content['summary']['timestamps']:
                    story.append(Paragraph('Important Timestamps', pdf_styles['heading']))
                    for timestamp in content['summary']['timestamps']:
                        story.append(Paragraph(f"• {timestamp}", pdf_styles['normal']))
                    story.append(Spacer(1, 20))
                
                # Key Takeaways
                if content['summary']['takeaways']:
                    story.append(Paragraph('Key Takeaways', pdf_styles['heading']))
                    for takeaway in content['summary']['takeaways']:
                        story.append(Paragraph(f"• {takeaway}", pdf_styles['normal']))
                    story.append(Spacer(1, 20))
                
                # Technical Information
                story.append(Paragraph('Technical Information', pdf_styles['heading']))
                
                tech_text = f"""
                <b>Transcript Available:</b> {'Yes' if content['transcript_info']['available'] else 'No'}<br/>
                """
                
                if content['transcript_info']['available']:
                    tech_text += f"""
                    <b>Language:</b> {content['transcript_info']['language']}<br/>
                    <b>Type:</b> {'Auto-generated' if content['transcript_info']['is_generated'] else 'Manual'}<br/>
                    """
                
                tech_text += f"""
                <b>AI Model:</b> {content['ai_info']['model']}<br/>
                <b>Tokens Used:</b> {content['ai_info']['tokens_used']}<br/>
                <b>Generated:</b> {content['ai_info']['generated_at']}
                """
                
                story.append(Paragraph(tech_text, pdf_styles['normal']))
                
                # Footer
                footer_text = f"Generated by {settings.bot_name} v{settings.bot_version}"
                self._add_pdf_footer(story, pdf_styles, footer_text)
                
            elif content_type == 'subtitles':
                # Video Information
                story.append(Paragraph('Информация о видео', pdf_styles['heading']))
                
                video_info_text = f"""
                <b>Название:</b> {content['video_info']['title']}<br/>
                <b>Канал:</b> {content['video_info']['channel']}<br/>
                <b>Длительность:</b> {content['video_info']['duration']}<br/>
                <b>ID видео:</b> {content['video_info']['video_id']}<br/>
                <b>Язык:</b> {content['video_info']['language']} ({content['video_info']['language_code']})<br/>
                <b>Тип:</b> {'Автогенерированные' if content['video_info']['auto_generated'] else 'Ручные'}<br/>
                <b>Количество сегментов:</b> {content['video_info']['subtitle_count']}
                """
                
                if content['video_info']['corrected']:
                    video_info_text += f"<br/><b>Обработка ИИ:</b> {content['video_info']['correction_method']}"
                
                story.append(Paragraph(video_info_text, pdf_styles['normal']))
                story.append(Spacer(1, 20))
                
                # Subtitles
                story.append(Paragraph('Субтитры', pdf_styles['heading']))
                
                # Split subtitles into paragraphs for better PDF formatting
                subtitle_lines = content['subtitles_text'].split('\n')
                for line in subtitle_lines:
                    if line.strip():
                        story.append(Paragraph(line.strip(), pdf_styles['normal']))
                
                # Footer
                footer_text = f"Создано: {content['generation_info']['generated_at']}<br/>"
                footer_text += f"Тип документа: {content['generation_info']['format']}<br/>"
                footer_text += f"Создано с помощью {settings.bot_name} v{settings.bot_version}"
                
                self._add_pdf_footer(story, pdf_styles, footer_text)
            
            # Build PDF
            doc.build(story)
        
        await self._run_in_executor(generate_pdf)
        return output_path

    async def create_subtitles_document(
        self, 
        subtitle_data: Dict, 
        output_format: str = 'txt'
    ) -> Optional[Path]:
        """
        Create formatted document with corrected subtitles.
        
        Args:
            subtitle_data: Corrected subtitle data from VideoSummarizer
            output_format: Output format ('txt', 'docx', 'pdf')
            
        Returns:
            Path to generated document or None on error
        """
        try:
            if output_format not in settings.supported_formats:
                raise ValueError(f"Unsupported format: {output_format}")

            # Prepare subtitle content
            content = self._prepare_subtitle_content(subtitle_data)
            
            # Generate filename
            video_title = subtitle_data.get('title', 'Unknown Video')
            corrected = subtitle_data.get('corrected', False)
            filename = self._generate_filename(video_title, 'subtitles', output_format, corrected)
            document_path = self.temp_dir / filename
            
            # Create document based on format
            if output_format == 'txt':
                document_path = await self._create_txt_document(content, document_path)
            elif output_format == 'docx':
                document_path = await self._create_docx_document(content, document_path)
            elif output_format == 'pdf':
                document_path = await self._create_pdf_document(content, document_path)
            
            logger.info(f"Created {output_format.upper()} subtitles document: {document_path}")
            return document_path
            
        except Exception as e:
            logger.error(f"Error creating subtitles document: {e}")
            return None
    
    def _prepare_subtitle_content(self, subtitle_data: Dict) -> Dict:
        """Prepare structured content for subtitle document generation."""
        # Format duration
        duration = subtitle_data.get('duration', 0)
        duration_str = f"{duration // 60}:{duration % 60:02d}"
        
        # Process subtitles
        subtitles = subtitle_data.get('subtitles', [])
        subtitle_text_parts = []
        
        for subtitle in subtitles:
            if hasattr(subtitle, 'text'):
                text = subtitle.text.strip()
                start_time = subtitle.start
            else:
                text = subtitle.get('text', '').strip()
                start_time = subtitle.get('start', 0)
            
            if text:
                # Format timestamp
                minutes = int(start_time // 60)
                seconds = int(start_time % 60)
                timestamp = f"[{minutes:02d}:{seconds:02d}]"
                subtitle_text_parts.append(f"{timestamp} {text}")
        
        return {
            'title': f"Субтитры: {subtitle_data.get('title', 'Unknown Title')}",
            'video_info': {
                'title': subtitle_data.get('title', 'Unknown Title'),
                'channel': subtitle_data.get('channel', 'Unknown'),
                'duration': duration_str,
                'video_id': subtitle_data.get('video_id', ''),
                'language': subtitle_data.get('language', 'Unknown'),
                'language_code': subtitle_data.get('language_code', ''),
                'auto_generated': subtitle_data.get('auto_generated', True),
                'corrected': subtitle_data.get('corrected', False),
                'correction_method': subtitle_data.get('correction_method', ''),
                'subtitle_count': len(subtitles)
            },
            'subtitles_text': '\n'.join(subtitle_text_parts),
            'generation_info': {
                'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC'),
                'format': 'Исправленные субтитры' if subtitle_data.get('corrected', False) else 'Субтитры'
            }
        }

    async def cleanup_old_documents(self, max_age_hours: int = 24):
        """Clean up old temporary documents."""
        try:
            import time
            current_time = time.time()
            cutoff_time = current_time - (max_age_hours * 3600)
            
            for file_path in self.temp_dir.glob('*'):
                if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                    try:
                        file_path.unlink()
                        logger.debug(f"Cleaned up old document: {file_path}")
                    except Exception as e:
                        logger.warning(f"Could not delete {file_path}: {e}")
                        
        except Exception as e:
            logger.error(f"Error during document cleanup: {e}")

    # Deprecated methods for backward compatibility
    async def _create_txt_subtitles_document(self, content: Dict, output_path: Path) -> Path:
        """Deprecated: Use _create_txt_document instead."""
        return await self._create_txt_document(content, output_path)

    async def _create_docx_subtitles_document(self, content: Dict, output_path: Path) -> Path:
        """Deprecated: Use _create_docx_document instead."""
        return await self._create_docx_document(content, output_path)

    async def _create_pdf_subtitles_document(self, content: Dict, output_path: Path) -> Path:
        """Deprecated: Use _create_pdf_document instead."""
        return await self._create_pdf_document(content, output_path) 